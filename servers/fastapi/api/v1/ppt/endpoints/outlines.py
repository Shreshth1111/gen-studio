from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from typing import Optional
from uuid import UUID
import asyncio
import io
import json

from database import get_session
from models.sql.presentation import Presentation
from models.sql.user import User
from api.v1.auth.endpoints import get_current_user
from services.llm_client import llm_client
from utils.streaming import StreamEvent, format_sse

router = APIRouter()

OUTLINE_SYSTEM = """You are an expert presentation architect.
Create a structured presentation outline in JSON format.
Respond ONLY with valid JSON:
{
  "title": "Presentation Title",
  "slides": [
    {"slide_number": 1, "title": "...", "layout_type": "title", "key_points": ["..."]}
  ]
}
Available layouts: title, bullets, two_column, image_left, image_right, stats, quote, timeline, comparison, team, blank, agenda"""


@router.get("/generate/stream")
async def stream_outline(
    presentation_id: UUID = Query(...),
    current_user: User = Depends(get_current_user),
    session: AsyncSession = Depends(get_session)
):
    result = await session.execute(
        select(Presentation).where(
            Presentation.id == presentation_id,
            Presentation.user_id == current_user.id
        )
    )
    pres = result.scalars().first()
    if not pres:
        raise HTTPException(404, "Presentation not found")

    events_queue: asyncio.Queue = asyncio.Queue()

    async def emit(event_type, data):
        await events_queue.put((event_type, data))

    async def generate_outline():
        try:
            await emit(StreamEvent.OUTLINE_START, {})
            prompt = f"""Topic: {pres.topic}
Audience: {pres.audience or 'professional'}
Tone: {pres.tone}
Language: {pres.language}
Slides: {pres.slide_count}
Density: {pres.content_density}

Create outline."""

            outline_text = ""
            async for token in llm_client.stream_text(OUTLINE_SYSTEM, prompt):
                outline_text += token
                await emit(StreamEvent.OUTLINE_CHUNK, {"token": token})

            try:
                cleaned = outline_text.strip()
                if cleaned.startswith("```"):
                    cleaned = cleaned.split("```")[1]
                    if cleaned.startswith("json"):
                        cleaned = cleaned[4:]
                data = json.loads(cleaned)
                slides = data.get("slides", [])
            except Exception:
                slides = [
                    {"slide_number": i+1, "title": f"Slide {i+1}",
                     "layout_type": "title" if i==0 else "bullets",
                     "key_points": [pres.topic]}
                    for i in range(pres.slide_count)
                ]

            await emit(StreamEvent.OUTLINE_DONE, {"outline": slides})

        except Exception as e:
            await emit(StreamEvent.ERROR, {"message": str(e)})
        finally:
            await events_queue.put(None)

    async def event_gen():
        task = asyncio.create_task(generate_outline())
        while True:
            item = await events_queue.get()
            if item is None:
                break
            et, d = item
            yield format_sse(et, d)
        await task

    return StreamingResponse(
        event_gen(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"}
    )


MAX_DOC_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_TEXT_CHARS = 30_000  # cap what we feed back to the LLM


def _extract_pdf(data: bytes) -> str:
    import pdfplumber
    out: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text:
                out.append(text)
    return "\n\n".join(out)


def _extract_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


@router.post("/parse-document")
async def parse_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    """Read a PDF or DOCX file and return its plain-text contents so the
    frontend can include it in the createPresentation payload as
    `source_text` for context-aware generation."""
    name = (file.filename or "").lower()
    if not (name.endswith(".pdf") or name.endswith(".docx")):
        raise HTTPException(400, "Only .pdf or .docx files are supported")

    data = await file.read()
    if len(data) > MAX_DOC_BYTES:
        raise HTTPException(413, f"File too large ({len(data)} bytes; max {MAX_DOC_BYTES})")

    try:
        text = (
            _extract_pdf(data) if name.endswith(".pdf") else _extract_docx(data)
        )
    except Exception as e:
        raise HTTPException(422, f"Failed to parse file: {e}")

    text = text.strip()
    if not text:
        raise HTTPException(422, "No readable text found in the document")

    truncated = len(text) > MAX_TEXT_CHARS
    if truncated:
        text = text[:MAX_TEXT_CHARS]

    return {
        "filename": file.filename,
        "length": len(text),
        "truncated": truncated,
        "text": text,
    }
